"""
Document generation service.
Produces DOCX files using python-docx from live ISMS data.
"""
import io
import json
from datetime import datetime

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

NAVY = RGBColor(0x1F, 0x38, 0x64)
BLUE = RGBColor(0x2E, 0x75, 0xB6)
GREEN = RGBColor(0x2C, 0x6E, 0x49)
LIGHT_BLUE = RGBColor(0xD6, 0xE4, 0xF0)


def _heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    run = p.runs[0] if p.runs else p.add_run(text)
    run.font.color.rgb = NAVY
    return p


def _add_footer(doc, org_name):
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.text = (f"Prepared in accordance with ISO/IEC 27001:2022 and the Zimbabwe Cyber and Data Protection Act (Chapter 12:07)\n"
              f"{org_name} | ISMS Compass | Generated {datetime.now().strftime('%d %b %Y')} | AI-Generated Draft — Requires Human Review and Approval before use as an official document")
    p.style = doc.styles['Footer']


def _add_document_header(doc, title_text, org):
    title = doc.add_heading(title_text, 0)
    title.runs[0].font.color.rgb = NAVY
    doc.add_paragraph(f"Organisation: {org.get('name', '')}")
    doc.add_paragraph(f"Sector: {org.get('sector', '')}")
    doc.add_paragraph(f"City: {org.get('city', '')}")
    doc.add_paragraph(f"Date Generated: {datetime.now().strftime('%d %B %Y')}")
    doc.add_paragraph()


def generate_information_security_policy(org: dict, session_data: dict) -> bytes:
    doc = Document()
    _add_footer(doc, org.get('name', 'Organisation'))

    # Title
    _add_document_header(doc, 'Information Security Policy', org)

    # Metadata table
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Table Grid'
    meta = [
        ('Organisation', org.get('name', '')),
        ('Version', '1.0'),
        ('Date', datetime.now().strftime('%d %B %Y')),
        ('Status', 'DRAFT — Awaiting Approval'),
    ]
    for i, (k, v) in enumerate(meta):
        table.rows[i].cells[0].text = k
        table.rows[i].cells[1].text = v

    doc.add_paragraph()

    _heading(doc, '1. Purpose and Scope', 1)
    doc.add_paragraph(
        f"This Information Security Policy establishes the management direction and commitment of "
        f"{org.get('name', 'the organisation')} to protecting its information assets in accordance with "
        f"ISO/IEC 27001:2022 and the Zimbabwe Cyber and Data Protection Act (Chapter 12:07)."
    )

    scope = session_data.get('scope_statement', f"This policy applies to all information systems, "
                             f"personnel, and third-party contractors of {org.get('name', 'the organisation')} "
                             f"operating from {org.get('city', 'Zimbabwe')}.")
    doc.add_paragraph(f"Scope: {scope}")

    _heading(doc, '2. Policy Statement', 1)
    policy_text = session_data.get('policy_text',
        f"{org.get('name', 'The organisation')} is committed to preserving the confidentiality, "
        f"integrity and availability of all information assets. Management recognises that information "
        f"security is fundamental to business operations and takes full accountability for the "
        f"implementation and maintenance of this Information Security Management System (ISMS).")
    doc.add_paragraph(policy_text)

    _heading(doc, '3. Information Security Objectives', 1)
    objectives = session_data.get('objectives', [
        'Achieve zero critical security incidents per quarter',
        'Maintain 100% staff information security training completion annually',
        'Implement 80% of applicable Annex A controls within 12 months',
    ])
    for i, obj in enumerate(objectives, 1):
        doc.add_paragraph(f"{i}. {obj}")

    _heading(doc, '4. Roles and Responsibilities', 1)
    doc.add_paragraph(
        "The ISMS Owner holds ultimate accountability for this policy. All staff are responsible "
        "for complying with this policy and reporting security incidents promptly. Specific roles "
        "and responsibilities are documented in the Roles & Responsibilities register (ISMS Step 7)."
    )

    _heading(doc, '5. Compliance', 1)
    doc.add_paragraph(
        "Non-compliance with this policy may result in disciplinary action. All breaches will be "
        "investigated and logged in the Audit Log. This policy will be reviewed annually or following "
        f"significant changes to the organisation or threat landscape."
    )

    _heading(doc, '6. Approval', 1)
    approved_by = session_data.get('signed_by', '_______________________')
    doc.add_paragraph(f"Approved by: {approved_by}")
    doc.add_paragraph(f"Date: {datetime.now().strftime('%d %B %Y')}")
    doc.add_paragraph("Signature: _______________________")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def generate_risk_assessment_report(org: dict, risks: list) -> bytes:
    doc = Document()
    _add_footer(doc, org.get('name', 'Organisation'))

    _add_document_header(doc, 'Risk Assessment Report', org)

    # Summary
    _heading(doc, '1. Executive Summary', 1)
    high = sum(1 for r in risks if r.get('risk_level') in ('High', 'Critical'))
    medium = sum(1 for r in risks if r.get('risk_level') == 'Medium')
    low = sum(1 for r in risks if r.get('risk_level') == 'Low')
    doc.add_paragraph(
        f"This Risk Assessment Report documents {len(risks)} identified information security risks "
        f"for {org.get('name', 'the organisation')}. Risk distribution: {high} High/Critical, "
        f"{medium} Medium, {low} Low. All risks have been assessed using the {org.get('risk_appetite', 'Standard')} "
        f"methodology in accordance with ISO 27001:2022 Clause 6.1.2."
    )

    _heading(doc, '2. Risk Register', 1)
    if risks:
        table = doc.add_table(rows=1, cols=7)
        table.style = 'Table Grid'
        headers = ['Risk ID', 'Asset', 'Threat', 'Likelihood', 'Impact', 'Score', 'Level']
        for i, h in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = h
            cell.paragraphs[0].runs[0].font.bold = True

        for risk in risks:
            row = table.add_row()
            vals = [
                risk.get('risk_id', ''),
                risk.get('asset', ''),
                (risk.get('threat', '') or '')[:50],
                str(risk.get('likelihood', '')),
                str(risk.get('impact', '')),
                str(risk.get('score', '')),
                risk.get('risk_level', ''),
            ]
            for i, v in enumerate(vals):
                row.cells[i].text = v
    else:
        doc.add_paragraph('No risks identified yet.')

    _heading(doc, '3. Risk Treatment Summary', 1)
    treatment_counts = {}
    for r in risks:
        t = r.get('treatment', 'Unknown')
        treatment_counts[t] = treatment_counts.get(t, 0) + 1
    for treatment, count in treatment_counts.items():
        doc.add_paragraph(f"• {treatment}: {count} risk(s)")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def generate_soa(org: dict, soa_entries: list) -> bytes:
    doc = Document()
    _add_footer(doc, org.get('name', 'Organisation'))

    _add_document_header(doc, 'Statement of Applicability', org)

    _heading(doc, '1. Introduction', 1)
    approved = sum(1 for e in soa_entries if e.get('human_approved'))
    doc.add_paragraph(
        f"This Statement of Applicability (SoA) documents the selection and justification of "
        f"ISO/IEC 27001:2022 Annex A controls for {org.get('name', 'the organisation')}. "
        f"{approved} of {len(soa_entries)} entries have been reviewed and approved. "
        f"Generated: {datetime.now().strftime('%d %B %Y')}."
    )

    _heading(doc, '2. Control Applicability', 1)
    if soa_entries:
        table = doc.add_table(rows=1, cols=6)
        table.style = 'Table Grid'
        headers = ['Ref', 'Control', 'Applicable', 'Justification', 'Status', 'Reviewed By']
        for i, h in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = h
            cell.paragraphs[0].runs[0].font.bold = True

        for entry in soa_entries:
            row = table.add_row()
            justification = entry.get('justification') or entry.get('ai_draft') or ''
            vals = [
                entry.get('annex_a_ref', ''),
                (entry.get('control_name', '') or '')[:40],
                'Yes' if entry.get('applicable') else 'No',
                justification[:80],
                entry.get('implementation_status', ''),
                entry.get('approved_by', ''),
            ]
            for i, v in enumerate(vals):
                row.cells[i].text = v
    else:
        doc.add_paragraph('No SoA entries yet. Complete the Statement of Applicability screen first.')

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def generate_risk_treatment_plan(org: dict, risks: list) -> bytes:
    doc = Document()
    _add_footer(doc, org.get('name', 'Organisation'))
    _add_document_header(doc, 'Risk Treatment Plan', org)

    _heading(doc, '1. Purpose', 1)
    doc.add_paragraph(
        f"This Risk Treatment Plan documents the treatment decisions for all identified information "
        f"security risks at {org.get('name', 'the organisation')} in accordance with ISO 27001:2022 Clause 6.1.3."
    )

    _heading(doc, '2. Treatment Register', 1)
    treated = [r for r in risks if r.get('treatment_plan')]
    if treated:
        for risk in treated:
            doc.add_heading(f"Risk {risk.get('risk_id', '')}: {risk.get('threat', '')}", 3)
            doc.add_paragraph(f"Treatment option: {risk.get('treatment', 'Mitigate')}")
            doc.add_paragraph(f"Control/Action: {risk.get('treatment_plan', 'TBD')}")
            doc.add_paragraph(f"Owner: {risk.get('treatment_owner', 'TBD')}")
            doc.add_paragraph(f"Target date: {risk.get('treatment_due', 'TBD')}")
    else:
        doc.add_paragraph('No treatment plans documented yet. Complete Step 5 first.')

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def generate_corrective_action_log(org: dict, actions: list) -> bytes:
    doc = Document()
    _add_footer(doc, org.get('name', 'Organisation'))
    _add_document_header(doc, 'Corrective Action Log', org)

    if actions:
        table = doc.add_table(rows=1, cols=6)
        table.style = 'Table Grid'
        for i, h in enumerate(['ID', 'Description', 'Source', 'Owner', 'Due Date', 'Status']):
            cell = table.rows[0].cells[i]
            cell.text = h
            cell.paragraphs[0].runs[0].font.bold = True
        for action in actions:
            row = table.add_row()
            for i, v in enumerate([
                action.get('action_id', ''), (action.get('description', '') or '')[:60],
                action.get('source', ''), action.get('assigned_to', ''),
                action.get('due_date', ''), action.get('status', ''),
            ]):
                row.cells[i].text = v
    else:
        doc.add_paragraph('No corrective actions logged yet.')

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def generate_scope_statement(org: dict, session_data: dict) -> bytes:
    doc = Document()
    _add_footer(doc, org.get('name', 'Organisation'))

    _add_document_header(doc, 'Scope Statement', org)

    _heading(doc, '1. Organisation Details', 1)
    doc.add_paragraph(f"Organisation Name: {org.get('name', '')}")
    doc.add_paragraph(f"Sector: {org.get('sector', '')}")
    doc.add_paragraph(f"City: {org.get('city', '')}")
    doc.add_paragraph(f"Risk Appetite: {org.get('risk_appetite', '')}")

    _heading(doc, '2. Scope Statement', 1)
    scope_statement = session_data.get('scopeStatement', '')
    if not scope_statement:
        scope_statement = session_data.get('scope_statement', 'Scope statement not defined.')
    doc.add_paragraph(scope_statement)

    _heading(doc, '3. Departments in Scope', 1)
    departments = session_data.get('departments', [])
    if departments and any(departments):
        for d in departments:
            if d:
                doc.add_paragraph(f"• {d}")
    else:
        doc.add_paragraph("No specific departments defined.")

    _heading(doc, '4. Locations in Scope', 1)
    locations = session_data.get('locations', [])
    if locations and any(locations):
        for l in locations:
            if l:
                doc.add_paragraph(f"• {l}")
    else:
        doc.add_paragraph("No specific locations defined.")

    _heading(doc, '5. Exclusions', 1)
    exclusions = session_data.get('exclusions', '')
    doc.add_paragraph(exclusions if exclusions else "None.")

    if 'applicable_legislation' in session_data:
        _heading(doc, '6. Applicable Legislation', 1)
        doc.add_paragraph(str(session_data['applicable_legislation']))

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


DOC_GENERATORS = {
    'information-security-policy': generate_information_security_policy,
    'risk-assessment-report': generate_risk_assessment_report,
    'soa': generate_soa,
    'risk-treatment-plan': generate_risk_treatment_plan,
    'corrective-action-log': generate_corrective_action_log,
    'scope-statement': generate_scope_statement,
}
